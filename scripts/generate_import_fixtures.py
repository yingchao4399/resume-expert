"""Generate deterministic Chinese import fixtures used by Playwright/manual QA."""

from pathlib import Path

from docx import Document
from reportlab.pdfbase import cidfonts, pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "e2e" / "fixtures"
OUTPUT.mkdir(parents=True, exist_ok=True)
LINES = [
    "张明 | 产品经理 | 上海",
    "电话：13800000000 | 邮箱：ming@example.com",
    "工作经历：示例科技，产品经理，2021 至今",
    "项目成果：主导库存流程重构，效率提升 40%",
]


def draw_lines(pdf: canvas.Canvas, font: str, page_count: int = 1) -> None:
    for page in range(page_count):
        pdf.setFont(font, 12)
        y = 800
        for repeat in range(12 if page_count > 1 else 1):
            for line in LINES:
                pdf.drawString(50, y, f"{line} 第 {page + 1} 页" if page_count > 1 else line)
                y -= 18
            if y < 80:
                break
        if page < page_count - 1:
            pdf.showPage()
    pdf.save()


pdfmetrics.registerFont(cidfonts.UnicodeCIDFont("STSong-Light"))
draw_lines(canvas.Canvas(str(OUTPUT / "chinese-cid-resume.pdf")), "STSong-Light")

pdfmetrics.registerFont(TTFont("SimHeiFixture", r"C:\Windows\Fonts\simhei.ttf"))
draw_lines(canvas.Canvas(str(OUTPUT / "chinese-ttf-resume.pdf")), "SimHeiFixture")
draw_lines(canvas.Canvas(str(OUTPUT / "multi-page-resume.pdf")), "SimHeiFixture", page_count=2)

document = Document()
document.add_heading("张明 - 产品经理", level=1)
for line in LINES[1:]:
    document.add_paragraph(line, style="List Bullet" if line.startswith("项目") else None)
document.save(OUTPUT / "chinese-resume.docx")

print(f"Generated fixtures in {OUTPUT}")
